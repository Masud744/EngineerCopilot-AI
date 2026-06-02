"""
EngineerCopilot AI — AI Resume Parsing Service.

Extracts text from PDF/DOCX and uses the LLM Abstraction Layer
to parse it into highly structured JSON data.
"""

from __future__ import annotations

import io
import json
import logging
import docx
import pdfplumber

from app.models.application import ResumeParseResponse
from app.ai.manager import get_ai_manager

logger = logging.getLogger(__name__)


def extract_pdf_text(file_bytes: bytes) -> str:
    """Extract raw text from PDF bytes."""
    text = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
    except Exception as e:
        logger.error("Error extracting text from PDF: %s", e)
    return "\n".join(text)


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract raw text from DOCX bytes."""
    text = []
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text.append(" | ".join(row_text))
    except Exception as e:
        logger.error("Error extracting text from DOCX: %s", e)
    return "\n".join(text)


async def parse_resume_file(file_bytes: bytes, file_ext: str) -> ResumeParseResponse:
    """
    Parse a resume using an LLM to extract structured fields.
    """
    # 1. Extract raw text
    if file_ext.lower() == "pdf":
        raw_text = extract_pdf_text(file_bytes)
    elif file_ext.lower() == "docx":
        raw_text = extract_docx_text(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")

    if not raw_text.strip():
        return ResumeParseResponse(raw_text="Empty resume file.")

    # 2. Setup LLM prompt for structured extraction
    system_prompt = """
    You are an expert ATS (Applicant Tracking System) parser.
    Your task is to analyze the provided resume text and extract the information into a strict JSON format.
    
    Output ONLY valid JSON. Do not include markdown blocks like ```json.
    
    Schema:
    {
      "skills": ["Python", "C++", "Machine Learning"],
      "education": [
        {
          "degree": "B.Sc Computer Science",
          "institution": "University Name",
          "start_date": "2018-01-01",
          "end_date": "2022-01-01",
          "description": "Any honors or relevant coursework"
        }
      ],
      "experience": [
        {
          "title": "Software Engineer",
          "company": "Tech Corp",
          "start_date": "2022-02-01",
          "end_date": "2024-01-01",
          "description": "Built scalable backend systems..."
        }
      ],
      "projects": [
        {
          "title": "Portfolio Website",
          "description": "Built using React and Next.js",
          "technologies": ["React", "Next.js"]
        }
      ],
      "certifications": [
        {
          "name": "AWS Certified Solutions Architect",
          "issuing_organization": "Amazon Web Services"
        }
      ]
    }
    
    Notes:
    - Format dates as YYYY-MM-DD. If only year is known, use YYYY-01-01. If current/present, leave end_date as null.
    - If a section is missing from the resume, return an empty array for that key.
    - Be as accurate as possible.
    """

    # 3. Call LLM Manager
    ai = get_ai_manager()
    try:
        response = await ai.complete(
            prompt=f"Resume Text:\n\n{raw_text[:8000]}", # limit text length for safety
            system_prompt=system_prompt,
            temperature=0.0, # low temp for JSON extraction
        )
        
        # 4. Clean and parse JSON response
        content = response.content.strip()
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        
        parsed_json = json.loads(content.strip())
        
        return ResumeParseResponse(
            skills=parsed_json.get("skills", []),
            education=parsed_json.get("education", []),
            experience=parsed_json.get("experience", []),
            projects=parsed_json.get("projects", []),
            certifications=parsed_json.get("certifications", []),
            raw_text=raw_text
        )

    except Exception as e:
        logger.error(f"AI parsing failed: {e}")
        # Return graceful degradation if AI fails
        return ResumeParseResponse(
            skills=[], education=[], experience=[], 
            projects=[], certifications=[], raw_text=raw_text
        )
